# api.py
# Windows 控制台编码修复
import sys
import io
if sys.platform == 'win32':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding='utf-8', errors='replace')

from fastapi import FastAPI, UploadFile, File, HTTPException, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.responses import HTMLResponse
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime, timedelta
import uvicorn
import re
import os
import shutil
import json 
import sqlite3
from contextlib import contextmanager
from dotenv import load_dotenv
from werkzeug.security import generate_password_hash, check_password_hash
from agent import chat_async, agent_executor
from PyPDF2 import PdfReader
# JWT 相关
from jose import JWTError, jwt

# 加载环境变量
load_dotenv()

# RAG 相关
from langchain_openai import ChatOpenAI
from langchain.prompts import PromptTemplate
from chunk import (
    load_vector_store, create_vector_store, add_new_documents_to_store,
    get_indexed_files, save_indexed_files, VECTOR_DB_DIR,
    get_chunked_files, save_chunked_files, load_cached_chunks, save_chunks_cache,
    BM25Retriever, HybridRetriever, BGEReranker,
    split_documents, chunk_only_new_files,
)
from hierarchical_splitter import HierarchicalTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader, Docx2txtLoader

# ==================== 配置 ====================
SECRET_KEY = os.getenv("JWT_SECRET_KEY", "change-this-in-production")
if SECRET_KEY == "change-this-in-production":
    print("[WARN] 警告：JWT_SECRET_KEY 使用默认值，请在 .env 文件中设置 JWT_SECRET_KEY")
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 60 * 24 * 7  # 7天

DOCUMENTS_DIR = "./documents"
_DOCS_DIR = DOCUMENTS_DIR
os.makedirs(DOCUMENTS_DIR, exist_ok=True)
os.makedirs(VECTOR_DB_DIR, exist_ok=True)

# ==================== 数据库 ====================
DB_PATH = "knowhub.db"

@contextmanager
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    try:
        yield conn
        conn.commit()
    finally:
        conn.close()

def init_db():
    with get_db() as conn:
        conn.execute("""
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                avatar TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS conversations (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                title TEXT DEFAULT '新会话',
                messages TEXT,
                files TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS user_files (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                filename TEXT NOT NULL,
                file_size INTEGER DEFAULT 0,
                page_count INTEGER DEFAULT 0,
                article_count INTEGER DEFAULT 0,
                word_count INTEGER DEFAULT 0,
                upload_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                chunk_count INTEGER DEFAULT 0,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS episodic_memory (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id TEXT NOT NULL,
                task_type TEXT NOT NULL,
                task_embedding TEXT NOT NULL,
                task_description TEXT,
                conversation_id INTEGER,
                mistake TEXT,
                fix TEXT,
                tool_trajectory TEXT,
                success_count INTEGER DEFAULT 0,
                fail_count INTEGER DEFAULT 0,
                hit_count INTEGER DEFAULT 1,
                last_matched_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("""
            CREATE TABLE IF NOT EXISTS dialogue_archive (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id TEXT NOT NULL,
                conversation_id INTEGER NOT NULL,
                turn INTEGER NOT NULL,
                role TEXT NOT NULL,
                content TEXT NOT NULL,
                content_embedding TEXT,
                tokens INTEGER DEFAULT 0,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("CREATE INDEX IF NOT EXISTS idx_dialogue_user ON dialogue_archive(user_id, created_at)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_dialogue_conv ON dialogue_archive(conversation_id, turn)")
        conn.execute("""
            CREATE TABLE IF NOT EXISTS knowledge_nuggets (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id TEXT NOT NULL,
                title TEXT NOT NULL,
                content TEXT NOT NULL,
                source_conv_id INTEGER,
                tags TEXT,
                content_embedding TEXT,
                confidence REAL DEFAULT 0.5,
                hit_count INTEGER DEFAULT 0,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """)
        conn.execute("CREATE INDEX IF NOT EXISTS idx_knowledge_user ON knowledge_nuggets(user_id, created_at)")
        # 如果表已存在但缺少新字段，添加它们
        try:
            conn.execute("ALTER TABLE user_files ADD COLUMN file_size INTEGER DEFAULT 0")
        except:
            pass
        try:
            conn.execute("ALTER TABLE user_files ADD COLUMN page_count INTEGER DEFAULT 0")
        except:
            pass
        try:
            conn.execute("ALTER TABLE user_files ADD COLUMN article_count INTEGER DEFAULT 0")
        except:
            pass
        try:
            conn.execute("ALTER TABLE user_files ADD COLUMN word_count INTEGER DEFAULT 0")
        except:
            pass
        # users 表添加 email 字段
        try:
            conn.execute("ALTER TABLE users ADD COLUMN email TEXT DEFAULT ''")
        except:
            pass
        # 订阅管理表
        conn.execute("""
            CREATE TABLE IF NOT EXISTS subscriptions (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                topic TEXT NOT NULL,
                frequency TEXT NOT NULL DEFAULT 'daily',
                email TEXT NOT NULL DEFAULT '',
                status TEXT DEFAULT 'active',
                last_run TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (user_id) REFERENCES users (id)
            )
        """)
        print("✅ 数据库初始化完成")

init_db()

# ==================== 密码加密 ====================

def get_password_hash(password):
    """生成密码哈希"""
    return generate_password_hash(password)

def verify_password(plain_password, hashed_password):
    """验证密码"""
    return check_password_hash(hashed_password, plain_password)

# ==================== JWT ====================
security = HTTPBearer()

def create_access_token(data: dict):
    to_encode = data.copy()
    expire = datetime.utcnow() + timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    token = credentials.credentials
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        return {"user_id": int(payload.get("sub")), "username": payload.get("username")}
    except JWTError:
        raise HTTPException(status_code=401, detail="无效的认证凭证")

# ==================== Pydantic 模型 ====================
class UserRegister(BaseModel):
    username: str
    password: str

class UserLogin(BaseModel):
    username: str
    password: str

class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"
    username: str
    user_id: int

class ConversationCreate(BaseModel):
    title: str = "新会话"

class ConversationUpdate(BaseModel):
    title: str
    messages: List[dict]
    files: List[dict]

class ChatRequest(BaseModel):
    message: str
    conversation_id: Optional[int] = None

class ChatResponse(BaseModel):
    answer: str
    conversation_id: int

class SubscriptionCreate(BaseModel):
    topic: str
    frequency: str = "daily"  # daily / weekly

class UserUpdate(BaseModel):
    email: str

# ==================== FastAPI 应用 ====================
app = FastAPI(title="知智 KnowHub API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ==================== LLM 实例 ====================
llm = ChatOpenAI(
    model="qwen3.7-plus",
    temperature=0.3,
    openai_api_key=os.getenv("DASHSCOPE_API_KEY"),
    openai_api_base="https://dashscope.aliyuncs.com/compatible-mode/v1"
)

prompt_template = """你是知智，企业级智能知识助手。请根据以下参考资料回答用户问题。

【对话历史】（如有）
{history}

【参考资料】
{context}

【用户问题】
{question}

【要求】
- 如果问题中有指代，参考对话历史理解
- 只根据参考资料回答，不确定就说"资料中没有相关信息"
- 引用具体的文档来源
- 用通俗语言解释专业术语
- 回答要完整、友好

【回答】"""

PROMPT = PromptTemplate(template=prompt_template, input_variables=["context", "question", "history"])

# ==================== RAG 检索器（懒初始化）====================
_hybrid_retriever = None
_reranker = None


def _init_rag():
    """启动时/上传文件后初始化混合检索器（只用缓存 chunks，不重新分块）"""
    global _hybrid_retriever, _reranker
    vs = load_vector_store()
    if vs is None:
        return False
    # 用 chunk_only_new_files:已有文件从缓存加载，只对新文件分块
    all_files = {f for f in os.listdir(_DOCS_DIR) if f.endswith(('.pdf', '.txt', '.docx', '.md'))}
    chunks = chunk_only_new_files(all_files, _DOCS_DIR)
    if not chunks:
        return False
    bm25 = BM25Retriever(chunks)
    _hybrid_retriever = HybridRetriever(vs, bm25)
    _reranker = BGEReranker()
    print(f"RAG 检索器就绪 ({len(chunks)} 个文本块)")
    return True


def _auto_extract_memory(user_id: str, question: str, answer: str):
    """自动从对话中提取用户的长期记忆（名字、偏好、身份等），存入 SQLite"""
    try:
        from tools import _get_memory_conn
        prompt = (
            f"用户的提问：{question[:300]}\n"
            f"助手的回答：{answer[:300]}\n\n"
            f"从以上对话中提取关于这个用户的重要信息。只提取明确的信息，不要推测。\n"
            f"提取项包括：用户名、职业、偏好风格、关注领域、重要上下文等。\n"
            f"如果没有可提取的新信息，回复 EMPTY。\n"
            f"如果有，用 JSON 格式回复：{{\"key\": \"value\", ...}}\n"
            f"只用中文回复，不要其他内容。"
        )
        resp = llm.invoke(prompt)
        text = resp.content.strip()
        if "EMPTY" in text or len(text) < 5:
            return

        # 尝试解析 JSON
        import re as _re
        json_match = _re.search(r'\{[^{}]+\}', text)
        if not json_match:
            return
        data = json.loads(json_match.group())

        conn = _get_memory_conn()
        for k, v in data.items():
            if isinstance(v, str) and len(v) >= 1 and len(k) >= 1:
                conn.execute(
                    "INSERT OR REPLACE INTO user_memory (user_id, memory_key, memory_value, updated_at) "
                    "VALUES (?, ?, ?, CURRENT_TIMESTAMP)",
                    (user_id, k.strip(), v.strip()[:500])
                )
        conn.commit()
        conn.close()
        if data:
            print(f"🧠 自动记忆提取: {list(data.keys())}")
    except Exception:
        pass  # 静默失败，不影响正常回答


# ==================== 辅助函数 ====================
def extract_file_metadata(filepath: str, text: str) -> dict:
    """提取文件的元数据（页数、条款数、字数等）"""
    metadata = {
        "file_size": os.path.getsize(filepath) // 1024,  # KB
        "page_count": 0,
        "article_count": 0,
        "word_count": len(text),
    }
    
    # 提取页数（如果是 PDF）
    if filepath.endswith('.pdf'):
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(filepath)
            metadata["page_count"] = len(reader.pages)
        except:
            metadata["page_count"] = 0
    elif filepath.endswith('.docx'):
        try:
            from docx import Document as DocxDocument
            docx_doc = DocxDocument(filepath)
            # docx 没有严格页数的概念，按段落数估算
            metadata["page_count"] = max(1, len(docx_doc.paragraphs) // 30)
        except:
            metadata["page_count"] = 0
    elif filepath.endswith('.md') or filepath.endswith('.txt'):
        metadata["page_count"] = max(1, metadata["word_count"] // 2000)
    
    # 提取条款数量（匹配"第X条"）
    import re
    article_pattern = r'第[一二三四五六七八九十\d]+条'
    articles = re.findall(article_pattern, text)
    metadata["article_count"] = len(set(articles))
    
    return metadata

# ==================== 认证接口 ====================
@app.post("/api/auth/register")
async def register(user: UserRegister):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM users WHERE username = ?", (user.username,))
        if cursor.fetchone():
            return {"code": 400, "message": "用户名已存在"}
        
        password_hash = get_password_hash(user.password)
        cursor.execute(
            "INSERT INTO users (username, password_hash) VALUES (?, ?)",
            (user.username, password_hash)
        )
        user_id = cursor.lastrowid
    
    return {"code": 200, "message": "注册成功", "data": {"user_id": user_id}}

@app.post("/api/auth/login")
async def login(user: UserLogin):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, username, password_hash, email FROM users WHERE username = ?", (user.username,))
        row = cursor.fetchone()
        if not row or not verify_password(user.password, row["password_hash"]):
            raise HTTPException(status_code=401, detail="用户名或密码错误")
        token = create_access_token({"sub": str(row["id"]), "username": row["username"]})
        return {"access_token": token, "token_type": "bearer",
                "username": row["username"], "user_id": row["id"],
                "email": row["email"] or ""}

@app.get("/api/auth/me")
async def get_current_user(user: dict = Depends(verify_token)):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT id, username, email, created_at FROM users WHERE id = ?", (user["user_id"],))
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="用户不存在")
        return {"code": 200, "data": dict(row)}

@app.put("/api/auth/me")
async def update_user_profile(data: UserUpdate, user: dict = Depends(verify_token)):
    """更新用户信息（邮箱等）"""
    with get_db() as conn:
        conn.execute(
            "UPDATE users SET email = ? WHERE id = ?",
            (data.email, user["user_id"])
        )
    return {"code": 200, "message": "更新成功"}

# ==================== 会话接口 ====================
@app.post("/api/conversations")
async def create_conversation(conv: ConversationCreate, user: dict = Depends(verify_token)):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "INSERT INTO conversations (user_id, title, messages, files) VALUES (?, ?, ?, ?)",
            (user["user_id"], conv.title, "[]", "[]")
        )
        conv_id = cursor.lastrowid
    return {"code": 200, "data": {"conversation_id": conv_id, "title": conv.title}}

@app.get("/api/conversations")
async def get_conversations(user: dict = Depends(verify_token)):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, title, created_at, updated_at FROM conversations WHERE user_id = ? ORDER BY updated_at DESC",
            (user["user_id"],)
        )
        rows = cursor.fetchall()
        conversations = []
        for row in rows:
            conv = dict(row)
            cursor.execute("SELECT messages FROM conversations WHERE id = ?", (conv["id"],))
            msg_row = cursor.fetchone()
            messages = json.loads(msg_row["messages"]) if msg_row["messages"] else []
            conv["message_count"] = len([m for m in messages if m.get("role") == "user"])
            conversations.append(conv)
    return {"code": 200, "data": conversations}

@app.get("/api/conversations/{conv_id}")
async def get_conversation(conv_id: int, user: dict = Depends(verify_token)):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, title, messages, files, created_at, updated_at FROM conversations WHERE id = ? AND user_id = ?",
            (conv_id, user["user_id"])
        )
        row = cursor.fetchone()
        if not row:
            raise HTTPException(status_code=404, detail="会话不存在")
        
        conv = dict(row)
        conv["messages"] = json.loads(conv["messages"]) if conv["messages"] else []
        conv["files"] = json.loads(conv["files"]) if conv["files"] else []
    return {"code": 200, "data": conv}

@app.put("/api/conversations/{conv_id}")
async def update_conversation(conv_id: int, conv: ConversationUpdate, user: dict = Depends(verify_token)):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "UPDATE conversations SET title = ?, messages = ?, files = ?, updated_at = CURRENT_TIMESTAMP WHERE id = ? AND user_id = ?",
            (conv.title, json.dumps(conv.messages, ensure_ascii=False), json.dumps(conv.files, ensure_ascii=False), conv_id, user["user_id"])
        )
        if cursor.rowcount == 0:
            raise HTTPException(status_code=404, detail="会话不存在")
    return {"code": 200, "message": "更新成功"}

@app.delete("/api/conversations/{conv_id}")
async def delete_conversation(conv_id: int, user: dict = Depends(verify_token)):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("DELETE FROM conversations WHERE id = ? AND user_id = ?", (conv_id, user["user_id"]))
    return {"code": 200, "message": "删除成功"}

# ==================== 文件上传接口 ====================
@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...), user: dict = Depends(verify_token)):
    filepath = os.path.join(DOCUMENTS_DIR, f"{user['user_id']}_{file.filename}")
    with open(filepath, "wb") as f:
        shutil.copyfileobj(file.file, f)
    
    # 解析文件
    if file.filename.endswith('.pdf'):
        loader = PyPDFLoader(filepath)
    elif file.filename.endswith('.txt') or file.filename.endswith('.md'):
        loader = TextLoader(filepath, encoding='utf-8')
    elif file.filename.endswith('.docx'):
        loader = Docx2txtLoader(filepath)
    else:
        return {"code": 400, "message": "暂不支持此格式"}
    
    docs = loader.load()
    for doc in docs:
        doc.metadata["source"] = file.filename
        doc.metadata["user_id"] = user["user_id"]
    
    # ========== 提取元数据 ==========
    full_text = "\n".join([d.page_content for d in docs])
    metadata = extract_file_metadata(filepath, full_text)
    print(f"📊 元数据: {metadata}")
    # ================================
    
    splitter = HierarchicalTextSplitter()
    chunks = splitter.split_documents(docs)
    
    vector_store = load_vector_store()
    if vector_store is None:
        vector_store = create_vector_store(chunks)
    else:
        vector_store = add_new_documents_to_store(vector_store, chunks)

    # 更新 chunk 缓存（追加新文件的 chunks，下次启动免分块）
    existing_chunks = load_cached_chunks()
    existing_chunks.extend(chunks)
    save_chunks_cache(existing_chunks)
    chunked_files = get_chunked_files()
    chunked_files.add(file.filename)
    save_chunked_files(chunked_files)
    
    # ========== 保存元数据到数据库 ==========
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute("""
            INSERT INTO user_files (user_id, filename, file_size, page_count, article_count, word_count, chunk_count) 
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """, (
            user["user_id"], 
            file.filename,
            metadata["file_size"],
            metadata["page_count"],
            metadata["article_count"],
            metadata["word_count"],
            len(chunks)
        ))
    # ======================================

    # 上传后重新初始化检索器
    _init_rag()

    return {"code": 200, "data": {
        "chunk_count": len(chunks), 
        "filename": file.filename,
        "metadata": metadata
    }}

@app.get("/api/files")
async def get_user_files(user: dict = Depends(verify_token)):
    with get_db() as conn:
        cursor = conn.cursor()
        cursor.execute(
            "SELECT id, filename, chunk_count, uploaded_at FROM user_files WHERE user_id = ? ORDER BY uploaded_at DESC",
            (user["user_id"],)
        )
        rows = cursor.fetchall()
        files = [dict(row) for row in rows]
    return {"code": 200, "data": files}

# ==================== 问答接口 ====================
def _save_chat(conv_id: int, user_id: int, messages: list, conv_title: str,
               question: str, answer: str) -> int:
    """保存消息到数据库，返回 conv_id"""
    messages.append({"role": "user", "content": question,
                     "timestamp": datetime.now().isoformat()})
    messages.append({"role": "assistant", "content": answer,
                     "timestamp": datetime.now().isoformat()})

    if conv_title == "新会话" and sum(1 for m in messages if m["role"] == "user") == 1:
        conv_title = question[:30] + ("..." if len(question) > 30 else "")

    with get_db() as conn:
        conn.execute(
            "UPDATE conversations SET title=?, messages=?, updated_at=CURRENT_TIMESTAMP WHERE id=?",
            (conv_title, json.dumps(messages, ensure_ascii=False), conv_id))
    return conv_id


@app.post("/api/chat", response_model=ChatResponse)
async def chat(request: ChatRequest, user: dict = Depends(verify_token)):
    try:
        # ---- 0. 会话管理 ----
        conv_id = request.conversation_id
        if conv_id is None:
            with get_db() as conn:
                c = conn.cursor()
                c.execute("INSERT INTO conversations (user_id,title,messages,files) VALUES (?,?,?,?)",
                          (user["user_id"], "新会话", "[]", "[]"))
                conv_id = c.lastrowid

        with get_db() as conn:
            row = conn.execute(
                "SELECT messages, title FROM conversations WHERE id=? AND user_id=?",
                (conv_id, user["user_id"])).fetchone()
            if not row:
                raise HTTPException(status_code=404, detail="会话不存在")
            messages = json.loads(row["messages"]) if row["messages"] else []
            conv_title = row["title"]

             # ========== 设置当前用户 ID（用于记忆隔离）==========
            from tools import set_current_user_id
            set_current_user_id(str(user["user_id"]))
        # ========== 👇 添加元数据查询 ==========
        # ---- 0.5 元数据/统计类问题（优先走元数据工具）----
        stats_keywords = ["多少条", "几条", "多少页", "几页", "多大", "文件大小", "总条数"]
        if any(kw in request.message for kw in stats_keywords):
            from tools import get_file_metadata
            import re
            # 提取文件名（匹配《XXX》或 XXX法/条例）
            match = re.search(r'《([^》]+)》', request.message)
            if not match:
                match = re.search(r'([\u4e00-\u9fa5]+(?:法|条例|办法|规定))', request.message)
            if match:
                filename = match.group(1)
                result = get_file_metadata.invoke({"filename": filename})
                if "未找到" not in result:
                    answer = result
                    conv_id = _save_chat(conv_id, user["user_id"], messages, conv_title,
                                         request.message, answer)
                    return ChatResponse(answer=answer, conversation_id=conv_id)

        # ========== LangGraph 多节点 Agent Workflow ==========
        # 构建对话历史
        history = []
        for msg in messages[-10:]:
            if msg["role"] == "user":
                history.append(("human", msg["content"]))
            else:
                history.append(("assistant", msg["content"]))

        # 调用 LangGraph 图（自动编排：记忆加载 → Query重写 → 意图澄清 → 意图路由 → RAG/Agent/闲聊 → 记忆保存）
        from agent import chat_async

        result = await chat_async(
            message=request.message,
            history=history,
            user_id=str(user["user_id"]),
            conversation_id=conv_id or 0,
        )

        answer = result["output"]

        # 保存聊天
        conv_id = _save_chat(conv_id, user["user_id"], messages, conv_title,
                             request.message, answer)

        return ChatResponse(answer=answer, conversation_id=conv_id)

    except Exception as e:
        import traceback as _tb
        print(f"Chat error: {e}")
        _tb.print_exc()
        return ChatResponse(answer=f"系统错误:{str(e)}",
                            conversation_id=request.conversation_id or 0)

# ==================== 公开 AI 问答接口（无需登录）====================
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id, filename, file_size, page_count, article_count,
                       word_count, chunk_count, upload_time
                FROM user_files
                ORDER BY upload_time DESC
            """)
            rows = cursor.fetchall()

        docs = []
        for row in rows:
            doc = dict(row)
            # 根据文件名推断分类
            filename = doc.get("filename", "")
            if any(kw in filename for kw in ["科技", "技术", "AI", "人工智能", "创新", "专利", "研发", "数字", "信息", "互联网", "智能"]):
                category = "科技"
            elif any(kw in filename for kw in ["税", "财务", "预算", "金融", "补贴", "价格", "贷款", "融资", "资产", "会计"]):
                category = "财税"
            elif any(kw in filename for kw in ["环保", "环境", "生态", "节能", "水", "污染", "流域", "气候", "气象", "碳", "绿色", "养殖", "灌溉"]):
                category = "环保"
            elif any(kw in filename for kw in ["规划", "纲要", "发展", "经济", "统计", "报告", "产业", "行业", "市场", "消费", "旅游", "景区"]):
                category = "行业报告"
            else:
                category = "政策法规"  # 默认分类

            docs.append({
                "id": doc["id"],
                "title": filename.replace(".pdf", "").replace(".txt", "").replace(".docx", "").replace(".md", ""),
                "summary": f"共{doc.get('page_count', 0)}页，{doc.get('word_count', 0)}字",
                "category": category,
                "file_type": ("pdf" if filename.endswith(".pdf") else
                              "docx" if filename.endswith(".docx") else
                              "md" if filename.endswith(".md") else "txt"),
                "created_at": doc.get("upload_time", ""),
                "updated_at": doc.get("upload_time", ""),
            })

        return docs
    except Exception as e:
        print(f"Get docs error: {e}")
        return []

@app.get("/api/docs/hot")
async def get_hot_docs():
    """获取热门文档"""
    docs = await get_docs()
    return docs[:6] if docs else []

@app.get("/api/docs/faq")
async def get_faq():
    """获取FAQ"""
    return [
        {"q": "如何使用AI助手？", "a": "点击页面右下角的AI助手图标，输入问题即可。"},
        {"q": "支持哪些文档格式？", "a": "目前支持PDF、Word（.docx）、Markdown（.md）和TXT格式的文档上传。"},
        {"q": "文档上传后多久可以使用？", "a": "文档上传后会自动解析和向量化，通常几秒到几分钟即可使用。"},
        {"q": "AI助手的回答准确吗？", "a": "AI基于知识库内容回答，准确率取决于知识库的质量。如果知识库没有相关内容，AI会自动联网搜索。"},
    ]

@app.get("/api/docs/catalog")
async def get_docs_catalog():
    """获取文档目录/分类列表"""
    try:
        docs = await get_docs()
        # 按分类组织文档
        catalog = {}
        for doc in docs:
            cat = doc.get("category", "其他")
            if cat not in catalog:
                catalog[cat] = {"category": cat, "items": []}
            catalog[cat]["items"].append(doc)

        return list(catalog.values())
    except Exception as e:
        print(f"Get catalog error: {e}")
        return []

@app.post("/api/docs")
async def create_doc(request: Request, user: dict = Depends(verify_token)):
    """创建新文档（管理后台）"""
    try:
        data = await request.json()
        title = data.get("title", "")
        category = data.get("category", "其他")
        content = data.get("content", "")

        if not title:
            raise HTTPException(status_code=400, detail="标题不能为空")

        # 创建文档文件
        filename = f"{title}.md"
        filepath = os.path.join(DOCUMENTS_DIR, filename)

        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)

        # 记录到数据库
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                INSERT INTO user_files (user_id, filename, file_size, word_count, upload_time)
                VALUES (?, ?, ?, ?, ?)
            """, (user["user_id"], filename, len(content.encode("utf-8")), len(content),
                  datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
            conn.commit()

        return {"status": "ok", "message": "文档创建成功"}
    except Exception as e:
        print(f"Create doc error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.put("/api/docs/{doc_id}")
async def update_doc(doc_id: int, request: Request, user: dict = Depends(verify_token)):
    """更新文档（管理后台）"""
    try:
        data = await request.json()
        title = data.get("title", "")
        category = data.get("category", "其他")
        content = data.get("content", "")

        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT filename FROM user_files WHERE id = ?", (doc_id,))
            row = cursor.fetchone()

            if not row:
                raise HTTPException(status_code=404, detail="文档不存在")

            old_filename = row["filename"]
            old_filepath = os.path.join(DOCUMENTS_DIR, old_filename)

            # 更新文件内容
            new_filename = f"{title}.md"
            new_filepath = os.path.join(DOCUMENTS_DIR, new_filename)

            with open(new_filepath, "w", encoding="utf-8") as f:
                f.write(content)

            # 如果文件名变了，删除旧文件
            if old_filename != new_filename and os.path.exists(old_filepath):
                os.remove(old_filepath)

            # 更新数据库
            cursor.execute("""
                UPDATE user_files
                SET filename = ?, word_count = ?, file_size = ?
                WHERE id = ?
            """, (new_filename, len(content), len(content.encode("utf-8")), doc_id))
            conn.commit()

        return {"status": "ok", "message": "文档更新成功"}
    except Exception as e:
        print(f"Update doc error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/admin/unanswered")
async def get_unanswered_questions():
    """获取未回答的问题列表（管理后台）"""
    try:
        # 从对话记录中找出AI回答为"未找到"或"无法回答"的问题
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id, user_id, title, messages, created_at
                FROM conversations
                ORDER BY created_at DESC
                LIMIT 100
            """)
            conversations = cursor.fetchall()

        unanswered = []
        for conv in conversations:
            messages = json.loads(conv["messages"]) if conv["messages"] else []
            for msg in messages:
                if msg.get("role") == "assistant":
                    content = msg.get("content", "")
                    # 检查是否是未回答或回答失败的情况
                    if any(keyword in content for keyword in ["未找到", "无法回答", "抱歉", "error", "错误"]):
                        # 找到对应的用户问题
                        user_msg = None
                        for prev_msg in reversed(messages[:messages.index(msg)]):
                            if prev_msg.get("role") == "user":
                                user_msg = prev_msg.get("content", "")
                                break

                        if user_msg:
                            unanswered.append({
                                "id": conv["id"],
                                "question": user_msg,
                                "answer": content,
                                "time": conv["created_at"]
                            })
                        break

        return unanswered[:20]  # 最多返回 20 条
    except Exception as e:
        print(f"Get unanswered error: {e}")
        return []

@app.get("/api/docs/{doc_id}")
async def get_doc_detail(doc_id: int):
    """获取文档详情，包含内容"""
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            cursor.execute("""
                SELECT id, filename, file_size, page_count, article_count,
                       word_count, chunk_count, upload_time
                FROM user_files WHERE id = ?
            """, (doc_id,))
            row = cursor.fetchone()

        if not row:
            return {"error": "文档不存在"}

        doc = dict(row)
        filename = doc.get("filename", "")

        # 根据文件名推断分类
        if any(kw in filename for kw in ["科技", "技术", "AI", "人工智能", "创新", "专利"]):
            category = "科技政策"
        elif any(kw in filename for kw in ["税", "财务", "预算", "金融", "补贴", "价格"]):
            category = "财税政策"
        elif any(kw in filename for kw in ["环保", "环境", "生态", "节能", "水", "污染", "流域"]):
            category = "环保政策"
        elif any(kw in filename for kw in ["规划", "纲要", "发展", "经济", "统计", "报告"]):
            category = "行业报告"
        else:
            category = "企业文档"

        # 动态计算页数和字数
        page_count = doc.get("page_count", 0)
        word_count = doc.get("word_count", 0)

        filepath = os.path.join(DOCUMENTS_DIR, filename)
        # 查找实际文件（可能带用户ID前缀）
        if not os.path.exists(filepath):
            for f in os.listdir(DOCUMENTS_DIR):
                if f.endswith(filename) or filename in f:
                    filepath = os.path.join(DOCUMENTS_DIR, f)
                    break

        if os.path.exists(filepath):
            if filepath.endswith('.txt') or filepath.endswith('.md'):
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        txt_content = f.read()
                    word_count = len(txt_content)
                    page_count = max(1, word_count // 2000)  # 按2000字/页估算
                except:
                    pass
            elif filepath.endswith('.pdf'):
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(filepath)
                    page_count = len(reader.pages)
                    # 统计字数
                    total_chars = 0
                    for page in reader.pages:
                        text = page.extract_text()
                        if text:
                            total_chars += len(text)
                    word_count = total_chars
                except:
                    pass
            elif filepath.endswith('.docx'):
                try:
                    from docx import Document as DocxDocument
                    docx_doc = DocxDocument(filepath)
                    total_chars = sum(len(p.text) for p in docx_doc.paragraphs)
                    word_count = total_chars
                    page_count = max(1, len(docx_doc.paragraphs) // 30)
                except:
                    pass

        if os.path.exists(filepath):
            if filepath.endswith('.txt') or filepath.endswith('.md'):
                try:
                    with open(filepath, 'r', encoding='utf-8') as f:
                        content = f.read()
                except:
                    with open(filepath, 'r', encoding='gbk', errors='ignore') as f:
                        content = f.read()
            elif filepath.endswith('.pdf'):
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(filepath)
                    pages_text = []
                    # 读取全部页面
                    for i, page in enumerate(reader.pages):
                        text = page.extract_text()
                        if text:
                            pages_text.append(f"## 第{i+1}页\n\n{text}")
                    content = "\n\n".join(pages_text)
                    if not content:
                        content = "*PDF内容提取失败，可能是扫描版PDF*"
                except Exception as e:
                    content = f"PDF解析失败: {str(e)}"
            elif filepath.endswith('.docx'):
                try:
                    from docx import Document as DocxDocument
                    docx_doc = DocxDocument(filepath)
                    paras = [p.text for p in docx_doc.paragraphs if p.text.strip()]
                    content = "\n\n".join(paras)
                    if not content:
                        content = "*Word文档内容提取失败*"
                except Exception as e:
                    content = f"Word文档解析失败: {str(e)}"

        # 生成Markdown格式的内容
        title = filename.replace(".pdf", "").replace(".txt", "").replace(".docx", "").replace(".md", "")
        markdown_content = f"""# {title}

**分类**: {category}
**页数**: {page_count} 页
**字数**: {word_count} 字
**上传时间**: {doc.get('upload_time', '')}

---

{content if content else '*文档内容暂无法预览，请联系管理员*'}
"""

        return {
            "id": doc["id"],
            "title": title,
            "summary": f"共{page_count}页，{word_count}字",
            "category": category,
            "content": markdown_content,
            "created_at": doc.get("upload_time", ""),
            "updated_at": doc.get("upload_time", ""),
        }
    except Exception as e:
        print(f"Get doc detail error: {e}")
        return {"error": str(e)}

@app.delete("/api/docs/{doc_id}")
async def delete_doc(doc_id: int, user: dict = Depends(verify_token)):
    """删除文档"""
    try:
        with get_db() as conn:
            cursor = conn.cursor()
            # 查询文档信息
            cursor.execute("SELECT filename FROM user_files WHERE id = ?", (doc_id,))
            row = cursor.fetchone()

            if not row:
                return {"error": "文档不存在"}

            filename = row[0]

            # 删除数据库记录
            cursor.execute("DELETE FROM user_files WHERE id = ?", (doc_id,))

            # 删除物理文件
            filepath = os.path.join(DOCUMENTS_DIR, filename)
            if os.path.exists(filepath):
                os.remove(filepath)
                print(f"已删除文件: {filename}")

            return {"message": "文档已删除"}
    except Exception as e:
        print(f"Delete doc error: {e}")
        return {"error": str(e)}

# ==================== 订阅管理接口 ====================

@app.get("/api/subscriptions")
async def list_subscriptions(user: dict = Depends(verify_token)):
    """列出当前用户的所有订阅"""
    with get_db() as conn:
        rows = conn.execute(
            "SELECT * FROM subscriptions WHERE user_id = ? ORDER BY created_at DESC",
            (user["user_id"],)
        ).fetchall()
        subs = [dict(r) for r in rows]
    return {"code": 200, "data": subs}


@app.post("/api/subscriptions")
async def create_subscription(data: SubscriptionCreate, user: dict = Depends(verify_token)):
    """创建新订阅"""
    # 获取用户邮箱
    with get_db() as conn:
        row = conn.execute(
            "SELECT email FROM users WHERE id = ?", (user["user_id"],)
        ).fetchone()
        email = row["email"] if row and row["email"] else ""

    if not email:
        return {"code": 400, "message": "请先在个人中心设置邮箱"}

    # 写入 subscriptions 表
    with get_db() as conn:
        cursor = conn.execute(
            "INSERT INTO subscriptions (user_id, topic, frequency, email) VALUES (?, ?, ?, ?)",
            (user["user_id"], data.topic, data.frequency, email)
        )
        sub_id = cursor.lastrowid

    # 同步到 scheduled_tasks.json（复用现有调度器）
    try:
        from tools import _load_tasks, _save_tasks, _get_next_id
        from datetime import timedelta
        tasks = _load_tasks()
        # 推送时间设为明天 09:00
        tomorrow = (datetime.now() + timedelta(days=1)).strftime("%Y-%m-%d")
        query = f"搜索{data.topic}领域最新发展动态和重要信息，整理成摘要后发送邮件到{email}，标题：{data.topic}推送"
        task = {
            "id": _get_next_id(tasks),
            "query": query,
            "time": f"{tomorrow} 09:00",
            "repeat": data.frequency,
            "created_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "last_run": None,
            "subscription_id": sub_id,  # 关联订阅 ID，方便取消时查找
        }
        tasks.append(task)
        _save_tasks(tasks)
    except Exception as e:
        print(f"[WARN] 同步到 scheduled_tasks.json 失败: {e}")

    return {"code": 200, "message": "订阅创建成功", "data": {"id": sub_id}}


@app.delete("/api/subscriptions/{sub_id}")
async def cancel_subscription(sub_id: int, user: dict = Depends(verify_token)):
    """取消订阅"""
    with get_db() as conn:
        # 确认是这个用户的订阅
        row = conn.execute(
            "SELECT id FROM subscriptions WHERE id = ? AND user_id = ?",
            (sub_id, user["user_id"])
        ).fetchone()
        if not row:
            return {"code": 404, "message": "订阅不存在"}
        conn.execute("DELETE FROM subscriptions WHERE id = ?", (sub_id,))

    # 从 scheduled_tasks.json 中移除对应任务
    try:
        from tools import _load_tasks, _save_tasks
        tasks = _load_tasks()
        tasks = [t for t in tasks if t.get("subscription_id") != sub_id]
        _save_tasks(tasks)
    except Exception as e:
        print(f"[WARN] 从 scheduled_tasks.json 移除失败: {e}")

    return {"code": 200, "message": "订阅已取消"}


# ==================== 健康检查 ====================
@app.get("/api/health")
async def health():
    return {"status": "ok"}

# ==================== 前端页面 ====================
@app.get("/", response_class=HTMLResponse)
async def root():
    html_content = '''<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>知智 KnowHub - 企业级智能知识助手</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap" rel="stylesheet">
    <link rel="stylesheet" href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.4.0/css/all.min.css">
    <script src="https://cdn.jsdelivr.net/npm/marked/marked.min.js"></script>
    <script src="https://cdn.tailwindcss.com"></script>
    <style>
        * { font-family: 'Inter', sans-serif; }
        ::-webkit-scrollbar { width: 6px; height: 6px; }
        ::-webkit-scrollbar-track { background: #f1f1f1; border-radius: 10px; }
        ::-webkit-scrollbar-thumb { background: #c1c1c1; border-radius: 10px; }
        .message-user { background: linear-gradient(135deg, #3b82f6, #2563eb); color: white; }
        .message-assistant { background-color: #f3f4f6; color: #1f2937; }
        .typing-indicator span { animation: blink 1.4s infinite; font-size: 20px; }
        .typing-indicator span:nth-child(2) { animation-delay: 0.2s; }
        .typing-indicator span:nth-child(3) { animation-delay: 0.4s; }
        @keyframes blink { 0%,60%,100% { opacity: 0.3; } 30% { opacity: 1; } }
        .upload-area { border: 2px dashed #e5e7eb; transition: all 0.2s ease; }
        .upload-area.drag-over { border-color: #3b82f6; background-color: #eff6ff; }
        .sidebar-item:hover { background-color: #f3f4f6; }
        .sidebar-item.active { background-color: #eff6ff; color: #3b82f6; }
        .history-item { cursor: pointer; transition: all 0.2s ease; }
        .history-item:hover { background-color: #f3f4f6; }
        .history-item.active { background-color: #eff6ff; border-left-color: #3b82f6; }
        .fade-in { animation: fadeIn 0.3s ease; }
        @keyframes fadeIn { from { opacity: 0; transform: translateY(10px); } to { opacity: 1; transform: translateY(0); } }
        .markdown-body { font-size: 14px; line-height: 1.6; }
        .markdown-body table { border-collapse: collapse; width: 100%; margin: 12px 0; }
        .markdown-body th, .markdown-body td { border: 1px solid #e5e7eb; padding: 8px 12px; }
        .markdown-body th { background-color: #f9fafb; font-weight: 600; }
        .markdown-body code { background-color: #f3f4f6; padding: 2px 6px; border-radius: 4px; }
        .recording { animation: pulse 1.5s infinite; background-color: #ef4444 !important; }
        @keyframes pulse { 0% { transform: scale(1); opacity: 1; } 50% { transform: scale(1.05); opacity: 0.8; } 100% { transform: scale(1); opacity: 1; } }
    </style>
</head>
<body class="bg-gray-50 h-screen overflow-hidden">

<div id="loginModal" class="fixed inset-0 bg-black bg-opacity-50 flex items-center justify-center z-50">
    <div class="bg-white rounded-2xl p-8 w-96 fade-in">
        <div class="text-center mb-6">
            <div class="w-16 h-16 bg-blue-500 rounded-2xl flex items-center justify-center mx-auto mb-4">
                <i class="fas fa-robot text-white text-3xl"></i>
            </div>
            <h2 class="text-2xl font-bold text-gray-800">知智 KnowHub</h2>
            <p class="text-gray-500 text-sm mt-1">企业级智能知识助手</p>
        </div>
        
        <div id="loginView">
            <div class="mb-4">
                <label class="block text-sm font-medium text-gray-700 mb-1">用户名</label>
                <input type="text" id="loginUsername" placeholder="请输入用户名" 
                       class="w-full border border-gray-200 rounded-lg px-4 py-2 focus:outline-none focus:ring-2 focus:ring-blue-500">
            </div>
            <div class="mb-4">
                <label class="block text-sm font-medium text-gray-700 mb-1">密码</label>
                <input type="password" id="loginPassword" placeholder="请输入密码" 
                       class="w-full border border-gray-200 rounded-lg px-4 py-2 focus:outline-none focus:ring-2 focus:ring-blue-500">
            </div>
            <button onclick="login()" class="w-full bg-blue-500 hover:bg-blue-600 text-white py-2 rounded-lg transition">登录</button>
            <p class="text-center text-sm text-gray-500 mt-4">新用户？<a href="javascript:void(0)" onclick="showRegister()" class="text-blue-500 hover:underline">立即注册</a></p>
        </div>
        
        <div id="registerView" class="hidden">
            <div class="mb-4">
                <label class="block text-sm font-medium text-gray-700 mb-1">用户名</label>
                <input type="text" id="regUsername" placeholder="请输入用户名" 
                       class="w-full border border-gray-200 rounded-lg px-4 py-2 focus:outline-none focus:ring-2 focus:ring-blue-500">
            </div>
            <div class="mb-4">
                <label class="block text-sm font-medium text-gray-700 mb-1">密码</label>
                <input type="password" id="regPassword" placeholder="请输入密码" 
                       class="w-full border border-gray-200 rounded-lg px-4 py-2 focus:outline-none focus:ring-2 focus:ring-blue-500">
            </div>
            <div class="mb-4">
                <label class="block text-sm font-medium text-gray-700 mb-1">确认密码</label>
                <input type="password" id="regConfirmPassword" placeholder="请再次输入密码" 
                       class="w-full border border-gray-200 rounded-lg px-4 py-2 focus:outline-none focus:ring-2 focus:ring-blue-500">
            </div>
            <button onclick="register()" class="w-full bg-blue-500 hover:bg-blue-600 text-white py-2 rounded-lg transition">注册</button>
            <p class="text-center text-sm text-gray-500 mt-4">已有账号？<a href="javascript:void(0)" onclick="showLogin()" class="text-blue-500 hover:underline">立即登录</a></p>
        </div>
    </div>
</div>

<div id="mainApp" class="hidden flex h-full">
    <aside class="w-80 bg-white border-r flex flex-col">
        <div class="p-5 border-b">
            <div class="flex items-center justify-between">
                <div class="flex items-center space-x-3">
                    <div class="w-10 h-10 bg-blue-500 rounded-xl flex items-center justify-center">
                        <i class="fas fa-robot text-white text-xl"></i>
                    </div>
                    <div><h1 class="text-xl font-bold text-gray-800">知智</h1><p class="text-xs text-gray-400">企业级智能知识助手</p></div>
                </div>
                <button onclick="logout()" class="text-gray-400 hover:text-gray-600"><i class="fas fa-sign-out-alt"></i></button>
            </div>
        </div>
        
        <div class="p-4 bg-gradient-to-r from-blue-50 to-indigo-50 border-b">
            <div class="flex items-center space-x-3">
                <div class="w-10 h-10 bg-blue-500 rounded-full flex items-center justify-center"><i class="fas fa-user text-white"></i></div>
                <div class="flex-1"><p class="text-sm font-semibold text-gray-800" id="currentUser">用户</p><p class="text-xs text-gray-500">普通账号</p></div>
                <div class="w-2 h-2 bg-green-500 rounded-full"></div>
            </div>
        </div>
        
        <div class="p-4"><button onclick="newConversation()" class="w-full bg-blue-500 hover:bg-blue-600 text-white py-2.5 rounded-xl flex items-center justify-center space-x-2 transition"><i class="fas fa-plus"></i><span>新建会话</span></button></div>
        
        <div class="px-4 pb-3">
            <div class="text-xs font-medium text-gray-500 mb-2"><i class="fas fa-cloud-upload-alt mr-1 text-blue-500"></i>上传文档</div>
            <div id="dropzone" class="upload-area bg-gray-50 rounded-xl p-4 text-center cursor-pointer">
                <i class="fas fa-file-upload text-gray-400 text-2xl mb-2"></i>
                <p class="text-xs text-gray-500">拖拽或点击上传</p>
                <p class="text-xs text-gray-400">支持 PDF / Word / Markdown / TXT</p>
                <input type="file" id="fileInput" class="hidden" accept=".pdf,.txt,.docx,.md" multiple>
            </div>
            <div id="file-list" class="mt-2 space-y-1 max-h-32 overflow-y-auto"><div class="text-xs text-gray-400 text-center py-2">暂无文档</div></div>
        </div>
        
        <div class="flex-1 overflow-y-auto px-3 pb-3">
            <div class="text-xs font-medium text-gray-500 mb-2 px-2"><i class="fas fa-history mr-1"></i>历史会话</div>
            <div id="historyList" class="space-y-1"><div class="text-xs text-gray-400 text-center py-4">加载中...</div></div>
        </div>
        
        <div class="p-4 border-t">
            <div class="flex justify-between text-xs text-gray-500">
                <span><i class="far fa-comment-dots mr-1"></i><span id="convCount">0</span> 次对话</span>
                <span><i class="far fa-file-alt mr-1"></i><span id="docCount">0</span> 个文档</span>
            </div>
        </div>
    </aside>

    <main class="flex-1 flex flex-col overflow-hidden">
        <div class="bg-white border-b px-6 py-3 flex items-center justify-between">
            <div class="flex items-center space-x-3"><i class="fas fa-comments text-blue-500"></i><h2 class="font-semibold text-gray-800" id="currentConvTitle">新会话</h2></div>
            <button onclick="clearCurrentChat()" class="text-gray-400 hover:text-gray-600 text-sm"><i class="fas fa-trash-alt"></i> 清空</button>
        </div>

        <div id="qa-messages" class="flex-1 overflow-y-auto p-6 space-y-4">
            <div class="flex justify-start fade-in">
                <div class="bg-gray-100 rounded-2xl rounded-tl-none px-4 py-3 max-w-2xl text-sm">
                    <div class="flex items-center space-x-2 mb-2"><i class="fas fa-robot text-blue-500"></i><span class="font-medium">知智</span></div>
                    您好！我是知智，企业级智能知识助手 🌟<br><br>
                    📎 左侧上传 PDF/Word/Markdown/TXT 文档，系统自动分块入库<br>
                    💬 然后输入问题，我会基于文档内容为您解答<br>
                    🎤 支持语音输入，点击麦克风按钮即可说话
                </div>
            </div>
        </div>

        <div class="p-4 bg-white border-t">
            <div class="flex space-x-2">
                <textarea id="qa-input" rows="1" placeholder="输入您的问题..." 
                          class="flex-1 border border-gray-200 rounded-xl px-4 py-3 text-sm focus:outline-none focus:ring-2 focus:ring-blue-500 resize-none"
                          style="max-height: 120px;"></textarea>
                <button id="voice-btn" class="bg-gray-100 hover:bg-gray-200 text-gray-600 px-4 py-3 rounded-xl transition"><i class="fas fa-microphone"></i></button>
                <button id="qa-send" class="bg-blue-500 hover:bg-blue-600 text-white px-6 py-3 rounded-xl transition"><i class="fas fa-paper-plane"></i></button>
            </div>
            <div class="flex justify-between text-xs text-gray-400 mt-2 px-1">
                <span><i class="far fa-keyboard"></i> Enter 发送，Shift+Enter 换行</span>
                <span><i class="fas fa-microphone"></i> 点击麦克风开始语音输入</span>
            </div>
        </div>
    </main>
</div>

<script>
    let token = null, currentUser = null, currentConvId = null, conversations = [], currentMessages = [], currentFiles = [];
    const API_BASE = '';
    let recognition = null, isRecording = false;
    
    if ('webkitSpeechRecognition' in window || 'SpeechRecognition' in window) {
        const SpeechRecognition = window.SpeechRecognition || window.webkitSpeechRecognition;
        recognition = new SpeechRecognition();
        recognition.continuous = false;
        recognition.interimResults = false;
        recognition.lang = 'zh-CN';
        recognition.onresult = (event) => { document.getElementById('qa-input').value = event.results[0][0].transcript; sendMessage(); };
        recognition.onerror = () => { addSystemMessage('语音识别失败'); stopRecording(); };
        recognition.onend = () => { stopRecording(); };
    }
    
    function startRecording() { if (recognition) { recognition.start(); isRecording = true; const btn = document.getElementById('voice-btn'); btn.classList.add('recording'); btn.innerHTML = '<i class="fas fa-microphone-slash"></i>'; } else { alert('您的浏览器不支持语音输入'); } }
    function stopRecording() { isRecording = false; const btn = document.getElementById('voice-btn'); btn.classList.remove('recording'); btn.innerHTML = '<i class="fas fa-microphone"></i>'; }
    
    async function fetchAPI(url, options = {}) {
        const headers = { 'Content-Type': 'application/json', ...options.headers };
        if (token) headers['Authorization'] = `Bearer ${token}`;
        const response = await fetch(`${API_BASE}${url}`, { ...options, headers });
        const data = await response.json();
        if (!response.ok) { if (response.status === 401) { logout(); throw new Error('登录已过期'); } throw new Error(data.detail || data.message); }
        return data;
    }
    
    async function register() {
        const username = document.getElementById('regUsername').value.trim();
        const password = document.getElementById('regPassword').value;
        const confirm = document.getElementById('regConfirmPassword').value;
        if (!username || !password) { alert('请填写用户名和密码'); return; }
        if (password !== confirm) { alert('两次输入的密码不一致'); return; }
        try {
            const data = await fetchAPI('/api/auth/register', { method: 'POST', body: JSON.stringify({ username, password }) });
            if (data.code === 200) { alert('注册成功！请登录'); showLogin(); } else { alert(data.message); }
        } catch(e) { alert(e.message); }
    }
    
    async function login() {
        const username = document.getElementById('loginUsername').value.trim();
        const password = document.getElementById('loginPassword').value;
        if (!username || !password) { alert('请输入用户名和密码'); return; }
        try {
            const response = await fetch(`${API_BASE}/api/auth/login`, { method: 'POST', headers: { 'Content-Type': 'application/json' }, body: JSON.stringify({ username, password }) });
            const data = await response.json();
            if (!response.ok) { alert(data.detail || '登录失败'); return; }
            token = data.access_token; currentUser = data.username;
            localStorage.setItem('knowhub_token', token); localStorage.setItem('knowhub_user', currentUser);
            document.getElementById('loginModal').classList.add('hidden');
            document.getElementById('mainApp').classList.remove('hidden');
            document.getElementById('currentUser').innerText = currentUser;
            await loadConversations();
            if (conversations.length === 0) await newConversation(); else await loadConversation(conversations[0].id);
            await loadFiles(); updateStats();
        } catch(e) { alert(e.message); }
    }
    
    async function logout() { token = null; currentUser = null; localStorage.removeItem('knowhub_token'); localStorage.removeItem('knowhub_user'); document.getElementById('loginModal').classList.remove('hidden'); document.getElementById('mainApp').classList.add('hidden'); }
    function showRegister() { document.getElementById('loginView').classList.add('hidden'); document.getElementById('registerView').classList.remove('hidden'); }
    function showLogin() { document.getElementById('registerView').classList.add('hidden'); document.getElementById('loginView').classList.remove('hidden'); }
    
    async function loadConversations() { const data = await fetchAPI('/api/conversations'); conversations = data.data || []; renderHistoryList(); updateStats(); }
    async function newConversation() { const data = await fetchAPI('/api/conversations', { method: 'POST', body: JSON.stringify({ title: '新会话' }) }); currentConvId = data.data.conversation_id; currentMessages = []; currentFiles = []; document.getElementById('currentConvTitle').innerText = '新会话'; renderMessages(); await loadConversations(); }
    async function loadConversation(convId) { currentConvId = convId; const data = await fetchAPI(`/api/conversations/${convId}`); const conv = data.data; currentMessages = conv.messages || []; currentFiles = conv.files || []; document.getElementById('currentConvTitle').innerText = conv.title || '会话'; renderMessages(); }
    async function deleteConversation(convId) { if (!confirm('确定删除？')) return; await fetchAPI(`/api/conversations/${convId}`, { method: 'DELETE' }); if (currentConvId === convId) { await loadConversations(); if (conversations.length > 0) await loadConversation(conversations[0].id); else await newConversation(); } else { await loadConversations(); } updateStats(); }
    async function clearCurrentChat() { currentMessages = []; await updateCurrentConversation(); renderMessages(); }
    async function updateCurrentConversation() { const conv = conversations.find(c => c.id === currentConvId); await fetchAPI(`/api/conversations/${currentConvId}`, { method: 'PUT', body: JSON.stringify({ title: conv?.title || '会话', messages: currentMessages, files: currentFiles }) }); }
    
    const dropzone = document.getElementById('dropzone'), fileInput = document.getElementById('fileInput');
    dropzone.onclick = () => fileInput.click();
    dropzone.ondragover = (e) => { e.preventDefault(); dropzone.classList.add('drag-over'); };
    dropzone.ondragleave = () => dropzone.classList.remove('drag-over');
    dropzone.ondrop = async (e) => { e.preventDefault(); dropzone.classList.remove('drag-over'); if (e.dataTransfer.files.length) { for (const f of e.dataTransfer.files) await uploadFile(f); } };
    fileInput.onchange = async (e) => { if (e.target.files.length) { for (const f of e.target.files) await uploadFile(f); } fileInput.value = ''; };
    
    async function uploadFile(file) {
        const formData = new FormData(); formData.append('file', file);
        addSystemMessage(`⏳ 正在上传「${file.name}」...`);
        try {
            const response = await fetch(`${API_BASE}/api/upload`, { method: 'POST', headers: { 'Authorization': `Bearer ${token}` }, body: formData });
            const data = await response.json();
            if (data.code === 200) { addSystemMessage(`✅ 已上传「${file.name}」，共 ${data.data.chunk_count} 个文本块`); await loadFiles(); updateStats(); }
            else { addSystemMessage(`❌ 上传失败:${data.message}`); }
        } catch(e) { addSystemMessage(`❌ 上传失败:${e.message}`); }
    }
    
    async function loadFiles() { const data = await fetchAPI('/api/files'); const files = data.data || []; const fileListDiv = document.getElementById('file-list'); if (files.length === 0) { fileListDiv.innerHTML = '<div class="text-xs text-gray-400 text-center py-2">暂无文档</div>'; return; } fileListDiv.innerHTML = files.map(f => `<div class="flex items-center justify-between text-xs py-1"><div class="flex items-center space-x-1"><i class="fas fa-file-pdf text-red-500 text-xs"></i><span class="text-gray-600 truncate max-w-36">${f.filename}</span></div><span class="text-gray-400 text-xs">${f.chunk_count || 0}块</span></div>`).join(''); }
    
    async function sendMessage() {
        const inputEl = document.getElementById('qa-input'); const question = inputEl.value.trim();
        if (!question) return;
        inputEl.value = ''; inputEl.style.height = 'auto';
        addMessage(question, true); addTypingIndicator();
        try {
            const data = await fetchAPI('/api/chat', { method: 'POST', body: JSON.stringify({ message: question, conversation_id: currentConvId }) });
            removeTypingIndicator(); addMessage(data.answer, false);
            currentConvId = data.conversation_id; await loadConversations(); await loadConversation(currentConvId); updateStats();
        } catch(e) { removeTypingIndicator(); addMessage(`发送失败:${e.message}`, false); }
    }
    
    function addMessage(content, isUser) { currentMessages.push({ role: isUser ? 'user' : 'assistant', content, timestamp: new Date().toISOString() }); renderMessages(); updateCurrentConversation(); }
    function addSystemMessage(content) { const div = document.createElement('div'); div.className = 'flex justify-center fade-in'; div.innerHTML = `<div class="text-xs text-gray-400 bg-gray-100 px-3 py-1 rounded-full">${content}</div>`; document.getElementById('qa-messages').appendChild(div); div.scrollIntoView({ behavior: 'smooth' }); setTimeout(() => div.remove(), 3000); }
    function addTypingIndicator() { const div = document.createElement('div'); div.id = 'typing'; div.className = 'flex justify-start fade-in'; div.innerHTML = '<div class="bg-gray-200 rounded-2xl rounded-tl-none px-4 py-2"><div class="typing-indicator flex space-x-1"><span>●</span><span>●</span><span>●</span></div></div>'; document.getElementById('qa-messages').appendChild(div); div.scrollIntoView({ behavior: 'smooth' }); }
    function removeTypingIndicator() { const el = document.getElementById('typing'); if (el) el.remove(); }
    function renderMessages() { const container = document.getElementById('qa-messages'); container.innerHTML = ''; currentMessages.forEach(msg => { const div = document.createElement('div'); div.className = `flex ${msg.role === 'user' ? 'justify-end' : 'justify-start'} fade-in`; const bubble = document.createElement('div'); bubble.className = `${msg.role === 'user' ? 'message-user' : 'message-assistant'} rounded-2xl ${msg.role === 'user' ? 'rounded-tr-none' : 'rounded-tl-none'} px-4 py-2 max-w-2xl text-sm`; if (msg.role === 'user') bubble.textContent = msg.content; else bubble.innerHTML = marked.parse(msg.content); div.appendChild(bubble); container.appendChild(div); }); container.scrollTop = container.scrollHeight; }
    function renderHistoryList() { const container = document.getElementById('historyList'); if (conversations.length === 0) { container.innerHTML = '<div class="text-xs text-gray-400 text-center py-4">暂无会话</div>'; return; } container.innerHTML = conversations.map(conv => `<div class="history-item p-2 rounded-lg ${currentConvId === conv.id ? 'active border-l-2 border-blue-500 bg-blue-50' : ''}" onclick="loadConversation(${conv.id})"><div class="flex items-center justify-between"><div class="flex-1"><div class="text-sm font-medium text-gray-700 truncate">${escapeHtml(conv.title || '新会话')}</div><div class="text-xs text-gray-400 mt-0.5">${conv.message_count || 0} 条消息</div></div><button onclick="event.stopPropagation(); deleteConversation(${conv.id})" class="text-gray-400 hover:text-red-500 text-xs"><i class="fas fa-trash-alt"></i></button></div></div>`).join(''); }
    function updateStats() { document.getElementById('convCount').innerText = conversations.length; loadFiles().catch(() => {}); }
    function escapeHtml(text) { const div = document.createElement('div'); div.textContent = text; return div.innerHTML; }
    
    document.getElementById('qa-send').onclick = sendMessage;
    document.getElementById('voice-btn').onclick = () => { if (isRecording) { if (recognition) recognition.stop(); stopRecording(); } else { startRecording(); } };
    document.getElementById('qa-input').onkeypress = (e) => { if (e.key === 'Enter' && !e.shiftKey) { e.preventDefault(); sendMessage(); } };
    document.getElementById('qa-input').addEventListener('input', function() { this.style.height = 'auto'; this.style.height = Math.min(this.scrollHeight, 120) + 'px'; });
    
    const savedToken = localStorage.getItem('knowhub_token');
    const savedUser = localStorage.getItem('knowhub_user');
    if (savedToken && savedUser) {
        token = savedToken; currentUser = savedUser;
        document.getElementById('loginModal').classList.add('hidden');
        document.getElementById('mainApp').classList.remove('hidden');
        document.getElementById('currentUser').innerText = currentUser;
        loadConversations().then(async () => { if (conversations.length > 0) await loadConversation(conversations[0].id); else await newConversation(); await loadFiles(); updateStats(); }).catch(() => logout());
    }
</script>
</body>
</html>'''
    return HTMLResponse(content=html_content)

# ==================== 后台定时任务调度器 ====================
_scheduler_thread = None

def _run_scheduler_loop():
    """后台线程：每60秒检查一次到期定时任务并自动执行"""
    import time as _time
    from tools import _load_tasks, _save_tasks
    from agent import agent_executor
    from datetime import datetime as _dt, timedelta as _td

    print("⏰ 定时任务调度器已启动（每60秒检查一次）")
    # 将日志写入文件方便调试
    import os as _os
    log_path = _os.path.join(_os.path.dirname(__file__), "scheduler.log")
    def _log(msg):
        try:
            with open(log_path, "a", encoding="utf-8") as lf:
                lf.write(f"{_dt.now().strftime('%H:%M:%S')} {msg}\n")
        except:
            pass
    _log("调度器启动")
    while True:
        try:
            tasks = _load_tasks()
            now = _dt.now()
            due = []
            for t in tasks:
                try:
                    exec_dt = _dt.strptime(t["time"], "%Y-%m-%d %H:%M")
                except ValueError:
                    continue
                if exec_dt <= now:
                    repeat = t.get("repeat", "once")
                    t["last_run"] = now.strftime("%Y-%m-%d %H:%M:%S")
                    if repeat == "once":
                        due.append(t)
                    elif repeat == "daily":
                        t["time"] = (exec_dt + _td(days=1)).strftime("%Y-%m-%d %H:%M")
                        due.append(t)
                    elif repeat == "weekly":
                        t["time"] = (exec_dt + _td(days=7)).strftime("%Y-%m-%d %H:%M")
                        due.append(t)

            # 移除已执行的一次性任务
            done_ids = {d["id"] for d in due}
            tasks = [t for t in tasks if t.get("repeat") != "once" or t["id"] not in done_ids]
            _save_tasks(tasks)

            # 执行到期任务
            for t in due:
                _log(f"执行任务 #{t['id']}: {t['query'][:80]}")
                try:
                    result = agent_executor.invoke({
                        "input": t["query"],
                        "chat_history": []
                    })
                    _log(f"任务 #{t['id']} 完成: {str(result.get('output', ''))[:120]}")
                    # 更新 subscriptions 表的 last_run
                    sub_id = t.get("subscription_id")
                    if sub_id:
                        try:
                            with get_db() as conn:
                                conn.execute(
                                    "UPDATE subscriptions SET last_run = ? WHERE id = ?",
                                    (now.strftime("%Y-%m-%d %H:%M:%S"), sub_id)
                                )
                        except:
                            pass
                except Exception as e:
                    _log(f"任务 #{t['id']} 失败: {e}")
        except Exception as e:
            _log(f"调度器异常: {e}")

        _time.sleep(60)


def start_scheduler():
    """启动后台定时任务调度器（幂等）"""
    global _scheduler_thread
    if _scheduler_thread and _scheduler_thread.is_alive():
        return
    import threading
    _scheduler_thread = threading.Thread(target=_run_scheduler_loop, daemon=True)
    _scheduler_thread.start()


# ==================== 启动 ====================
if __name__ == "__main__":
    start_scheduler()
    uvicorn.run(app, host="0.0.0.0", port=8000)